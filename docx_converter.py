import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def json_to_docx(json_data, output_filename="esports_report.docx"):
    if isinstance(json_data, str):
        try:
            data = json.loads(json_data)
        except json.JSONDecodeError:
            raise ValueError("Invalid JSON string provided")
    else:
        data = json_data
    
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    for block in data:
        if not isinstance(block, dict):
            continue
            
        block_type = block.get("type")
        
        if block_type == "heading":
            try:
                level = block.get("level", 1)
                text = block.get("text", "")
                
                clean_text = clean_markdown(text)
                
                heading = doc.add_heading(clean_text, level=level)
                
                for run in heading.runs:
                    if level == 1:
                        run.font.color.rgb = RGBColor(0, 0, 128)
                        run.font.size = Pt(16)
                    else:
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        run.font.size = Pt(14 - (level - 2))
            except Exception:
                doc.add_heading("Heading", level=1)
        
        elif block_type == "paragraph":
            try:
                text = block.get("text", "")
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                process_text_with_markdown(paragraph, text)
            except Exception:
                doc.add_paragraph()
            
        elif block_type == "list":
            try:
                items = block.get("items", [])
                for item in items:
                    paragraph = doc.add_paragraph(style='List Bullet')
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
                    if isinstance(item, dict):
                        if 'text' in item:
                            process_text_with_markdown(paragraph, item['text'])
                        else:
                            process_text_with_markdown(paragraph, str(item))
                    elif isinstance(item, list):
                        joined_text = ', '.join(str(x) for x in item)
                        process_text_with_markdown(paragraph, joined_text)
                    else:
                        process_text_with_markdown(paragraph, item)
            except Exception:
                doc.add_paragraph("• List item", style='List Bullet')
    
    try:
        doc.save(output_filename)
        return output_filename
    except Exception:
        fallback_filename = "report_fallback.docx"
        doc.save(fallback_filename)
        return fallback_filename


def clean_markdown(text):
    if not isinstance(text, str):
        try:
            text = str(text)
        except:
            return ""
            
    try:
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'(?<!\w)\*(?!\*)(.*?)(?<!\*)\*(?!\w)', r'\1', text)
        text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
        text = re.sub(r'`(.*?)`', r'\1', text)
        
        return text
    except Exception:
        return text


def process_text_with_markdown(paragraph, text):
    try:
        if not isinstance(text, str):
            try:
                text = str(text)
            except:
                text = ""
                
        bold_pattern = re.compile(r'\*\*(.*?)\*\*')
        
        positions = []
        for match in bold_pattern.finditer(text):
            positions.append((match.start(), 'start_bold'))
            positions.append((match.end(), 'end_bold'))
        
        if not positions:
            paragraph.add_run(text)
            return
        
        positions.sort()
        
        last_pos = 0
        is_bold = False
        
        for pos, marker_type in positions:
            if pos > last_pos:
                run_text = text[last_pos:pos]
                if is_bold and run_text.startswith('**'):
                    run_text = run_text[2:]
                if not is_bold and run_text.endswith('**'):
                    run_text = run_text[:-2]
                    
                if run_text:
                    run = paragraph.add_run(run_text)
                    run.bold = is_bold
            
            if marker_type == 'start_bold':
                is_bold = True
            elif marker_type == 'end_bold':
                is_bold = False
            
            last_pos = pos
        
        if last_pos < len(text):
            run = paragraph.add_run(text[last_pos:])
            run.bold = is_bold
    except Exception:
        try:
            paragraph.add_run(str(text))
        except:
            paragraph.add_run("Text processing error")


def main():
    import argparse
    
    parser = argparse.ArgumentParser(description='Convert JSON to Word document')
    parser.add_argument('--input', '-i', type=str, help='Input JSON file path')
    parser.add_argument('--output', '-o', type=str, default='esports_report.docx', 
                        help='Output DOCX file path')
    parser.add_argument('--test', action='store_true', help='Run with test data')
    args = parser.parse_args()
    
    if args.test:
        test_data = [
            {"type": "heading", "level": 1, "text": "Test Document"},
            {"type": "paragraph", "text": "This is a paragraph with **bold text** inside it."},
            {"type": "heading", "level": 2, "text": "List Items"},
            {"type": "list", "items": [
                "Regular item",
                "Item with **bold text**",
                ["This", "is", "a", "list", "item"],
                {"text": "This is a dict item", "value": 100},
                "**Completely bold item**"
            ]}
        ]
        output_path = json_to_docx(test_data, args.output)
        print(f"Test document created: {output_path}")
        return
    
    try:
        if args.input:
            with open(args.input, 'r') as f:
                json_data = json.load(f)
        else:
            print("No input file specified. Using data from main.py output.")
            try:
                from main import report_response
                json_text = report_response.text
                json_text = re.sub(r'^```json\s*', '', json_text)
                json_text = re.sub(r'\s*```$', '', json_text)
                json_data = json.loads(json_text)
            except ImportError:
                print("Couldn't import from main.py. Using sample data instead.")
                json_data = [
                    {"type": "heading", "level": 1, "text": "Sample Report"},
                    {"type": "paragraph", "text": "This is a sample report."}
                ]
        
        output_path = json_to_docx(json_data, args.output)
        print(f"Document created successfully: {output_path}")
    except Exception as e:
        print(f"An error occurred: {e}")
        print("Creating fallback document...")
        fallback_data = [
            {"type": "heading", "level": 1, "text": "Error Report"},
            {"type": "paragraph", "text": f"An error occurred during processing: {str(e)}"}
        ]
        output_path = json_to_docx(fallback_data, "error_report.docx")
        print(f"Fallback document created: {output_path}")


if __name__ == "__main__":
    main() 